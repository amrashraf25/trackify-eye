"""
template_processor.py
---------------------
Replaces {{Placeholder}} tokens inside a Word (.docx) template with values
from a data dict, then saves the result to a specified output path.

Handles:
  • Paragraphs (including placeholders split across multiple runs)
  • Tables (all cells)
  • Headers and footers
  • Dynamic columns — any key present in the data dict is substituted
"""

import logging
import re
from pathlib import Path
from typing import Any, Dict, Set

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# Matches  {{AnyWord}}  or  {{Any Phrase}}
PLACEHOLDER_RE = re.compile(r"\{\{([^}]+)\}\}")


# ── Low-level helpers ────────────────────────────────────────────────────────

def _replace_in_paragraph(paragraph, data: Dict[str, Any]) -> None:
    """
    Replace all {{key}} occurrences in a single paragraph.

    Word sometimes splits a placeholder across several Run objects (e.g. when
    the user typed it with autocorrect active, or changed formatting mid-token).
    We detect this by checking the *full* paragraph text, and when a placeholder
    is present we consolidate all runs into one before doing the substitution so
    that formatting of the first run is preserved.
    """
    # Fast path: nothing to do
    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return

    # If the placeholder is cleanly inside a single run we can replace it there.
    # Otherwise we fall back to the merge strategy.
    needs_merge = False
    for run in paragraph.runs:
        if PLACEHOLDER_RE.search(run.text):
            run.text = _substitute(run.text, data)

    # Re-check full text for any cross-run placeholders
    full_text_after = "".join(run.text for run in paragraph.runs)
    if PLACEHOLDER_RE.search(full_text_after):
        needs_merge = True

    if needs_merge and paragraph.runs:
        # Merge all run text into the first run, clear the rest
        merged = "".join(run.text for run in paragraph.runs)
        merged = _substitute(merged, data)
        paragraph.runs[0].text = merged
        for run in paragraph.runs[1:]:
            run.text = ""


def _substitute(text: str, data: Dict[str, Any]) -> str:
    """Replace every {{key}} in *text* with data[key] (if key exists)."""
    def replacer(match: re.Match) -> str:
        key = match.group(1).strip()
        return str(data[key]) if key in data else match.group(0)
    return PLACEHOLDER_RE.sub(replacer, text)


def _process_paragraphs(paragraphs, data: Dict[str, Any]) -> None:
    for para in paragraphs:
        _replace_in_paragraph(para, data)


def _process_table(table, data: Dict[str, Any]) -> None:
    for row in table.rows:
        for cell in row.cells:
            _process_paragraphs(cell.paragraphs, data)
            # Nested tables
            for nested in cell.tables:
                _process_table(nested, data)


# ── Public API ───────────────────────────────────────────────────────────────

def get_placeholders(template_path: str) -> Set[str]:
    """
    Scan a template and return all unique placeholder names found (without
    the {{ }} delimiters).
    """
    doc = Document(template_path)
    found: Set[str] = set()

    def _scan(text: str) -> None:
        for m in PLACEHOLDER_RE.finditer(text):
            found.add(m.group(1).strip())

    for para in doc.paragraphs:
        _scan(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan(para.text)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _scan(para.text)
        for para in section.footer.paragraphs:
            _scan(para.text)

    return found


def process_template(
    template_path: str,
    data: Dict[str, Any],
    output_path: str,
) -> None:
    """
    Open *template_path*, substitute all {{key}} tokens using *data*, and
    write the result to *output_path*.

    Parameters
    ----------
    template_path : Path to the .docx template file.
    data          : Mapping of placeholder names → replacement values.
    output_path   : Destination path for the generated .docx file.
    """
    if not Path(template_path).exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    # Body paragraphs
    _process_paragraphs(doc.paragraphs, data)

    # Tables
    for table in doc.tables:
        _process_table(table, data)

    # Headers & footers
    for section in doc.sections:
        _process_paragraphs(section.header.paragraphs, data)
        _process_paragraphs(section.footer.paragraphs, data)

    # Ensure output directory exists
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info("Generated: %s", output_path)


def safe_filename(value: str, fallback: str = "document") -> str:
    """
    Convert an arbitrary string into a safe Windows/macOS/Linux filename.
    Strips characters that are illegal in filenames on any major OS.
    """
    # Remove characters illegal on Windows (also covers Linux/macOS)
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", value).strip(". ")
    return cleaned if cleaned else fallback
