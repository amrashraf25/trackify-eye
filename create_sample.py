"""
create_sample.py
----------------
Run this script ONCE to create a ready-to-use sample:
  • data.xlsx     — three rows of student data
  • template.docx — a certificate template with {{Name}}, {{Course}},
                    {{Grade}}, and {{Date}} placeholders

Usage
-----
    python create_sample.py
"""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


# ── Excel ─────────────────────────────────────────────────────────────────────

def create_excel(path: str = "data.xlsx") -> None:
    data = {
        "Name":   ["Alice Johnson", "Bob Smith",    "Carol White"],
        "Course": ["Python 101",    "Data Science", "Web Development"],
        "Grade":  ["A",             "B+",            "A-"],
        "Date":   ["2024-01-15",    "2024-01-16",   "2024-01-17"],
    }
    df = pd.DataFrame(data)
    df.to_excel(path, index=False)
    print(f"  Created: {path}")


# ── Word template ─────────────────────────────────────────────────────────────

def create_template(path: str = "template.docx") -> None:
    doc = Document()

    # ── Title
    title = doc.add_heading("Certificate of Completion", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()   # spacer

    # ── Body
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run("This is to certify that").font.size = Pt(12)

    doc.add_paragraph()

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run("{{Name}}")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()

    course_intro = doc.add_paragraph()
    course_intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course_intro.add_run("has successfully completed the course:").font.size = Pt(12)

    doc.add_paragraph()

    course_para = doc.add_paragraph()
    course_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = course_para.add_run("{{Course}}")
    run2.bold = True
    run2.font.size = Pt(16)

    doc.add_paragraph()

    # ── Grade & Date
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run("Grade: ").bold = True
    details.add_run("{{Grade}}")
    details.add_run("     Date: ").bold = True
    details.add_run("{{Date}}")

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Footer text
    footer = doc.add_paragraph("Congratulations on your achievement!")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.save(path)
    print(f"  Created: {path}")


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Creating sample files …")
    create_excel()
    create_template()
    print(
        "\nAll done! Now run:\n"
        "    python main.py\n\n"
        "Click 'Generate Now' and check the 'output' folder."
    )
